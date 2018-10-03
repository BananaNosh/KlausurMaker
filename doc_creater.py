from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
import re


def create_klausur_template(filename, title, sentences, sentence_delimiters, vocabs_to_show, indices_of_shown,
                            show_numbers=False):
    document = Document("./templates/default.docx")
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(14)

    heading = document.add_heading("")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    font = heading.add_run(title).font
    font.name = "Times New Roman"
    font.size = Pt(18)
    font.color.rgb = RGBColor(0, 0, 0)

    sub_heading = document.add_heading("", level=2)
    sub_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub_heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = sub_heading.add_run("Thema der Unterrichtseinheit: Bestimmt was von Cicero")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)

    list_para = document.add_paragraph('Übersetze bitte ...', style='List Number')
    list_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    list_para.style.font.bold = True

    body = document.add_paragraph("")
    body.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    word_regex = re.compile(r"\b[^\d\W]+\b")
    for i, (sentence, indices_to_underline) in enumerate(zip(sentences, indices_of_shown)):
        indices_to_underline = list(set(indices_to_underline))
        indices_to_underline.sort()
        match = word_regex.search(sentence)
        word_index = 0
        not_underlined_start = 0
        for index in indices_to_underline:
            while word_index < index:
                match = word_regex.search(sentence, pos=match.span()[1])
                word_index += 1
            match_end = match.span()[1]
            match_start = match.span()[0]
            body.add_run(sentence[not_underlined_start:match_start])
            body.add_run(sentence[match_start:match_end]).underline = True
            not_underlined_start = match_end
        body.add_run(sentence[not_underlined_start:])
        if i < len(sentences) - 1:
            body.add_run(sentence_delimiters[i])

    help_style = document.styles.add_style("Vocab Help", WD_STYLE_TYPE.PARAGRAPH)
    help_style.base_style = document.styles["Normal"]
    help_style.font.size = Pt(9)
    help_style.font.bold = True
    vocab_help_para = document.add_paragraph("", style="Vocab Help")
    vocab_help_para.add_run("Übersetzungshilfen:").italic = True
    for i, sentence in enumerate(vocabs_to_show):
        if show_numbers and len(sentence) > 0:
            vocab_help_para.add_run(f" {i+1}")
        for j, (word, lemma, adds, translation) in enumerate(sentence):
            adds_string = f", {adds}" if adds is not None and len(adds) > 0 else ""
            adds_string = adds_string.replace("-", u"\u2011")
            vocab_string = f" {lemma}{adds_string} \u2011 {translation}"
            vocab_string = (" " if j != 0 else u"\xa0") + vocab_string.replace(" ", u"\xa0").strip()
            vocab_help_para.add_run(vocab_string)
            vocab_help_para.add_run("," if j < len(sentence) - 1 else ".")

    list_para = document.add_paragraph('Aufgaben zur Interretation:', style='List Number')
    list_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    document.save(filename)


if __name__ == '__main__':
    filename = "demo.docx"
    title = "1. Lateinklausur E1 26.09.18"
    sentences = [
        '1 Quare secedant improbi, secernant se a bonis; in unum locum congregentur, muro denique secernantur a nobis.',
        '2 Desinant insidari consuli, obsidere curiam cum gladiis, comparare faces ad imflammandam urbem.',
        '3 Denique in fronte uniuscuiusque inscriptum sit, quid de re publica sentiat.',
        '4 Tu, Catilina, ad salutem rei publicae, ad tuam pestem et perniciem, ad exitium eorum, qui se tecum omni '
        'scelere iunxerunt, proficiscere ad impium et nefarium bellum!',
        '5 Tu, Juppiter, inimicos bonorum, hostes patriae, '
        'latrones Italiae aeternis suppliciis vivos mortuosque mactabis.']
    vocabs = [[('Quare', 'quare', '', 'Deshalb'), ('congregentur', 'congregari', 'congregatus sum', 'sich zusammenscharen')],
              [('Desinant', 'desinare', '', 'aufhören'),
               ('insidari', 'insidari', 'c. Dat', 'jemandem nach dem Leben trachten'),
               ('faces', 'fax', 'facis f', 'Fackel')], [], [], [("latrones", "latro", "-onis m", "Räuber")]]
    create_klausur_template(filename, title, sentences, [" ", " ", "\n", "\n"],
                            vocabs, [[0, 10], [0, 1, 8], [], [], []], True)
